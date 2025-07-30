"""测试新增5个功能的综合测试"""

import os
import tempfile
import unittest
from pathlib import Path

# 添加src目录到路径
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

from word_mcp.create_document import create_document
from word_mcp.add_paragraph import add_paragraph
from word_mcp.add_heading import add_heading
from word_mcp.insert_header_near_text import insert_header_near_text
from word_mcp.insert_line_or_paragraph_near_text import insert_line_or_paragraph_near_text
from word_mcp.add_picture import add_picture
from word_mcp.get_document_text import get_document_text
from word_mcp.exceptions import FileError, DocumentError, ValidationError
from docx import Document


class TestNewBatchFeatures(unittest.TestCase):
    """测试新增5个功能的综合测试类"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.test_file = os.path.join(self.temp_dir, "test_document.docx")

    def tearDown(self):
        """清理测试环境"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_add_paragraph_basic(self):
        """测试基本的添加段落功能"""
        create_document(filepath=self.test_file)

        result = add_paragraph(
            filepath=self.test_file,
            text="这是一个测试段落"
        )

        self.assertIn("成功向文档添加段落", result["message"])
        self.assertEqual(result["style"], "Normal")

    def test_add_heading_basic(self):
        """测试基本的添加标题功能"""
        create_document(filepath=self.test_file)

        result = add_heading(
            filepath=self.test_file,
            text="测试标题",
            level=1
        )

        self.assertIn("成功向文档添加1级标题", result["message"])
        self.assertEqual(result["level"], 1)

    def test_insert_header_near_text_basic(self):
        """测试基本的在指定文本附近插入标题功能"""
        create_document(filepath=self.test_file)

        # 先添加一个段落作为目标
        add_paragraph(filepath=self.test_file, text="目标段落内容")

        result = insert_header_near_text(
            filepath=self.test_file,
            target_text="目标段落",
            header_title="插入的标题"
        )

        self.assertTrue(result["found"])
        self.assertEqual(result["header_title"], "插入的标题")

    def test_insert_line_near_text_basic(self):
        """测试基本的在指定文本附近插入段落功能"""
        create_document(filepath=self.test_file)

        # 先添加一个段落作为目标
        add_paragraph(filepath=self.test_file, text="目标段落内容")

        result = insert_line_or_paragraph_near_text(
            filepath=self.test_file,
            target_text="目标段落",
            line_text="插入的新段落"
        )

        self.assertTrue(result["found"])
        self.assertEqual(result["line_text"], "插入的新段落")

    def test_add_picture_missing_image(self):
        """测试添加不存在的图片（应该失败）"""
        create_document(filepath=self.test_file)

        fake_image = os.path.join(self.temp_dir, "fake.png")

        with self.assertRaises(FileError) as context:
            add_picture(
                filepath=self.test_file,
                image_path=fake_image
            )
        self.assertIn("图片文件不存在", str(context.exception))

    def test_comprehensive_workflow(self):
        """测试综合工作流程"""
        # 1. 创建文档
        create_document(filepath=self.test_file, title="测试文档")

        # 2. 添加标题
        add_heading(filepath=self.test_file, text="第一章 概述", level=1)

        # 3. 添加段落
        add_paragraph(filepath=self.test_file, text="这是第一章的内容。")

        # 4. 在指定文本后插入标题
        insert_header_near_text(
            filepath=self.test_file,
            target_text="第一章",
            header_title="1.1 子标题",
            position="after"
        )

        # 5. 在指定文本后插入段落
        insert_line_or_paragraph_near_text(
            filepath=self.test_file,
            target_text="这是第一章的内容",
            line_text="这是补充说明。",
            position="after"
        )

        # 验证内容
        doc_text = get_document_text(self.test_file)
        self.assertIn("第一章 概述", doc_text["text"])
        self.assertIn("1.1 子标题", doc_text["text"])
        self.assertIn("这是第一章的内容", doc_text["text"])
        self.assertIn("这是补充说明", doc_text["text"])

    def test_error_handling(self):
        """测试错误处理"""
        # 文件不存在的情况
        non_existent = os.path.join(self.temp_dir, "non_existent.docx")

        with self.assertRaises(FileError):
            add_paragraph(filepath=non_existent, text="测试")

        with self.assertRaises(FileError):
            add_heading(filepath=non_existent, text="测试")

        # 空参数的情况
        create_document(filepath=self.test_file)

        with self.assertRaises(ValidationError):
            add_paragraph(filepath=self.test_file, text="")

        with self.assertRaises(ValidationError):
            add_heading(filepath=self.test_file, text="")

    def test_heading_levels(self):
        """测试不同级别的标题"""
        create_document(filepath=self.test_file)

        for level in [1, 2, 3, 4, 5]:
            result = add_heading(
                filepath=self.test_file,
                text=f"级别 {level} 标题",
                level=level
            )
            self.assertEqual(result["level"], level)

        # 测试无效级别
        with self.assertRaises(ValidationError):
            add_heading(filepath=self.test_file, text="测试", level=10)


if __name__ == '__main__':
    unittest.main()