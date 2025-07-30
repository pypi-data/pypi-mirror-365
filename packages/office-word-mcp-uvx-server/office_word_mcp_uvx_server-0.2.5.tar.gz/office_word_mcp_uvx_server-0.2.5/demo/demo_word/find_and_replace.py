
import re
from docx import Document

def replace_placeholder_in_cell(cell, placeholder, replacement):
    """
    替换单元格中的占位符。
    """
    for p in cell.paragraphs:
        # 将一个段落的runs拼接起来，处理分散的占位符
        full_text = ''.join(run.text for run in p.runs)
        if placeholder in full_text:
            # 清除旧内容并添加新内容
            for run in p.runs:
                run.text = ''
            p.runs[0].text = full_text.replace(placeholder, replacement)

def find_and_replace_in_docx(file_path, replacements):
    """
    查找并替换Word文档中的占位符。
    """
    doc = Document(file_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, replacement in replacements.items():
                    # 我们需要重新组合文本，因为一个单元格的文本可能被分割
                    if placeholder in ''.join(run.text for run in cell.paragraphs[0].runs):
                         replace_placeholder_in_cell(cell, placeholder, replacement)

    # 保存修改后的文档
    # 注意：这里我们直接修改了原始文件，如果需要，可以先复制一份
    new_file_path = file_path.replace('.docx', '_modified.docx')
    doc.save(new_file_path)
    print(f"已将修改后的文件保存为: {new_file_path}")

if __name__ == "__main__":
    file_to_modify = "xxx.docx"
    
    # 定义要替换的内容
    replacements = {
        "%%2": "283±1",
        # 在这里添加更多的替换规则
    }

    find_and_replace_in_docx(file_to_modify, replacements)
