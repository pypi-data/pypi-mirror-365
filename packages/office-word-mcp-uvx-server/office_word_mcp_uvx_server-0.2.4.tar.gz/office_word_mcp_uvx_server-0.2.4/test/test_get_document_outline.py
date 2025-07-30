"""测试get_document_outline功能的单元测试"""

import os
import tempfile
import unittest
from pathlib import Path

# 添加src目录到路径
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

from word_mcp.create_document import create_document
from word_mcp.get_document_outline import get_document_outline
from word_mcp.exceptions import FileError, DocumentError, ValidationError
from docx import Document


class TestGetDocumentOutline(unittest.TestCase):
    """测试get_document_outline功能的单元测试类"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.test_file = os.path.join(self.temp_dir, "test_document.docx")

    def tearDown(self):
        """清理测试环境"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_get_document_outline_empty_document(self):
        """测试获取空文档的大纲"""
        # 创建空文档
        create_document(filepath=self.test_file)

        # 获取文档大纲
        result = get_document_outline(filepath=self.test_file)

        self.assertIn("outline", result)
        self.assertIn("summary", result)
        self.assertEqual(len(result["outline"]["paragraphs"]), 0)
        self.assertEqual(len(result["outline"]["tables"]), 0)
        self.assertEqual(len(result["outline"]["headings"]), 0)
        self.assertEqual(result["summary"]["total_paragraphs"], 0)
        self.assertEqual(result["summary"]["table_count"], 0)
        self.assertEqual(result["summary"]["heading_count"], 0)

    def test_get_document_outline_with_paragraphs(self):
        """测试获取包含段落的文档大纲"""
        # 创建文档并添加段落
        create_document(filepath=self.test_file)

        doc = Document(self.test_file)
        doc.add_paragraph("这是普通段落。")
        doc.add_paragraph("这是另一个普通段落，内容更长一些，用来测试文本预览功能。")
        doc.save(self.test_file)

        # 获取文档大纲
        result = get_document_outline(filepath=self.test_file)

        outline = result["outline"]
        summary = result["summary"]

        self.assertEqual(len(outline["paragraphs"]), 2)
        self.assertEqual(summary["total_paragraphs"], 2)
        self.assertEqual(summary["normal_paragraph_count"], 2)
        self.assertEqual(summary["heading_count"], 0)

        # 检查段落信息
        para1 = outline["paragraphs"][0]
        self.assertEqual(para1["index"], 0)
        self.assertEqual(para1["text_preview"], "这是普通段落。")
        self.assertFalse(para1["is_heading"])

    def test_get_document_outline_with_headings(self):
        """测试获取包含标题的文档大纲"""
        create_document(filepath=self.test_file)

        doc = Document(self.test_file)
        doc.add_heading("主标题", level=1)
        doc.add_paragraph("正文段落")
        doc.add_heading("副标题", level=2)
        doc.add_paragraph("更多正文内容")
        doc.save(self.test_file)

        # 获取文档大纲
        result = get_document_outline(filepath=self.test_file)

        outline = result["outline"]
        summary = result["summary"]

        self.assertEqual(len(outline["headings"]), 2)
        self.assertEqual(summary["heading_count"], 2)
        self.assertEqual(summary["normal_paragraph_count"], 2)
        self.assertTrue(summary["has_structured_headings"])

        # 检查标题信息
        heading1 = outline["headings"][0]
        self.assertEqual(heading1["text"], "主标题")
        self.assertEqual(heading1["level"], 1)

        heading2 = outline["headings"][1]
        self.assertEqual(heading2["text"], "副标题")
        self.assertEqual(heading2["level"], 2)

    def test_get_document_outline_with_tables(self):
        """测试获取包含表格的文档大纲"""
        create_document(filepath=self.test_file)

        doc = Document(self.test_file)
        doc.add_paragraph("文档开始")

        # 添加小表格
        table1 = doc.add_table(rows=2, cols=2)
        table1.cell(0, 0).text = "表1-行1列1"
        table1.cell(0, 1).text = "表1-行1列2"
        table1.cell(1, 0).text = "表1-行2列1"
        table1.cell(1, 1).text = "表1-行2列2"

        doc.add_paragraph("中间段落")

        # 添加大表格
        table2 = doc.add_table(rows=4, cols=3)
        table2.cell(0, 0).text = "大表格标题行第一列"
        table2.cell(0, 1).text = "大表格标题行第二列"
        table2.cell(0, 2).text = "大表格标题行第三列"

        doc.save(self.test_file)

        # 获取文档大纲
        result = get_document_outline(filepath=self.test_file)

        outline = result["outline"]
        summary = result["summary"]

        self.assertEqual(len(outline["tables"]), 2)
        self.assertEqual(summary["table_count"], 2)
        self.assertTrue(summary["has_tables"])

        # 检查第一个表格
        table1_info = outline["tables"][0]
        self.assertEqual(table1_info["rows"], 2)
        self.assertEqual(table1_info["columns"], 2)
        self.assertEqual(len(table1_info["preview_data"]), 2)

        # 检查第二个表格
        table2_info = outline["tables"][1]
        self.assertEqual(table2_info["rows"], 4)
        self.assertEqual(table2_info["columns"], 3)
        self.assertEqual(len(table2_info["preview_data"]), 3)  # 最多3行预览

    def test_get_document_outline_file_not_exists(self):
        """测试获取不存在文件的大纲应该失败"""
        non_existent_file = os.path.join(self.temp_dir, "non_existent.docx")

        with self.assertRaises(FileError) as context:
            get_document_outline(non_existent_file)
        self.assertIn("文件不存在", str(context.exception))

    def test_get_document_outline_invalid_format(self):
        """测试获取非docx文件的大纲应该失败"""
        txt_file = os.path.join(self.temp_dir, "test.txt")
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write("这不是docx文件")

        with self.assertRaises(FileError) as context:
            get_document_outline(txt_file)
        self.assertIn("文件格式不支持", str(context.exception))

    def test_get_document_outline_empty_filepath(self):
        """测试空文件路径应该失败"""
        with self.assertRaises(ValidationError) as context:
            get_document_outline(filepath="")
        self.assertIn("文件路径不能为空", str(context.exception))

    def test_get_document_outline_return_fields(self):
        """测试返回结果包含所有必需的字段"""
        # 创建文档
        create_document(filepath=self.test_file)

        # 获取文档大纲
        result = get_document_outline(filepath=self.test_file)

        # 检查顶级字段
        top_level_fields = ["message", "file_path", "outline", "summary"]
        for field in top_level_fields:
            self.assertIn(field, result, f"缺少顶级字段: {field}")

        # 检查outline字段
        outline_fields = ["paragraphs", "tables", "headings"]
        for field in outline_fields:
            self.assertIn(field, result["outline"], f"outline缺少字段: {field}")

        # 检查summary字段
        summary_fields = [
            "total_paragraphs", "heading_count", "normal_paragraph_count",
            "table_count", "has_structured_headings", "has_tables", "document_sections"
        ]
        for field in summary_fields:
            self.assertIn(field, result["summary"], f"summary缺少字段: {field}")

    def test_get_document_outline_text_preview_truncation(self):
        """测试长文本的预览截断功能"""
        create_document(filepath=self.test_file)

        doc = Document(self.test_file)
        long_text = "这是一个非常长的段落，" * 20  # 创建超过100字符的文本
        doc.add_paragraph(long_text)
        doc.save(self.test_file)

        # 获取文档大纲
        result = get_document_outline(filepath=self.test_file)

        para_info = result["outline"]["paragraphs"][0]
        preview = para_info["text_preview"]

        # 检查文本是否被正确截断
        self.assertLessEqual(len(preview), 103)  # 100 + "..." = 103
        self.assertTrue(preview.endswith("..."))


if __name__ == '__main__':
    unittest.main()