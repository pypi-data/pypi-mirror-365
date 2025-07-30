"""测试get_document_text功能的单元测试"""

import os
import tempfile
import unittest
from pathlib import Path

# 添加src目录到路径
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

from word_mcp.create_document import create_document
from word_mcp.get_document_text import get_document_text
from word_mcp.exceptions import FileError, DocumentError, ValidationError
from docx import Document


class TestGetDocumentText(unittest.TestCase):
    """测试get_document_text功能的单元测试类"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.test_file = os.path.join(self.temp_dir, "test_document.docx")

    def tearDown(self):
        """清理测试环境"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_get_document_text_empty_document(self):
        """测试获取空文档的文本"""
        # 创建空文档
        create_document(filepath=self.test_file)

        # 获取文档文本
        result = get_document_text(filepath=self.test_file)

        self.assertEqual(result["text"], "")
        self.assertEqual(result["paragraph_count"], 0)
        self.assertEqual(result["table_count"], 0)
        self.assertEqual(result["total_characters"], 0)
        self.assertEqual(result["total_words"], 0)

    def test_get_document_text_with_paragraphs(self):
        """测试获取包含段落的文档文本"""
        # 创建文档并添加段落
        create_document(filepath=self.test_file)

        # 直接使用docx库添加内容用于测试
        doc = Document(self.test_file)
        doc.add_paragraph("这是第一个段落。")
        doc.add_paragraph("这是第二个段落，包含更多文字。")
        doc.add_paragraph("")  # 空段落应该被忽略
        doc.add_paragraph("第三个段落。")
        doc.save(self.test_file)

        # 获取文档文本
        result = get_document_text(filepath=self.test_file)

        expected_text = "这是第一个段落。\n这是第二个段落，包含更多文字。\n第三个段落。"
        self.assertEqual(result["text"], expected_text)
        self.assertEqual(result["paragraph_count"], 3)  # 空段落被忽略
        self.assertEqual(result["table_count"], 0)
        self.assertGreater(result["total_characters"], 0)
        self.assertGreater(result["total_words"], 0)

    def test_get_document_text_with_table(self):
        """测试获取包含表格的文档文本"""
        # 创建文档并添加表格
        create_document(filepath=self.test_file)

        doc = Document(self.test_file)
        doc.add_paragraph("文档段落内容")

        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "表头1"
        table.cell(0, 1).text = "表头2"
        table.cell(1, 0).text = "数据1"
        table.cell(1, 1).text = "数据2"
        doc.save(self.test_file)

        # 获取文档文本
        result = get_document_text(filepath=self.test_file)

        self.assertIn("文档段落内容", result["text"])
        self.assertIn("--- 表格内容 ---", result["text"])
        self.assertIn("表头1", result["text"])
        self.assertIn("表头2", result["text"])
        self.assertIn("数据1", result["text"])
        self.assertIn("数据2", result["text"])
        self.assertEqual(result["paragraph_count"], 1)
        self.assertEqual(result["table_count"], 1)

    def test_get_document_text_file_not_exists(self):
        """测试获取不存在文件的文本应该失败"""
        non_existent_file = os.path.join(self.temp_dir, "non_existent.docx")

        with self.assertRaises(FileError) as context:
            get_document_text(non_existent_file)
        self.assertIn("文件不存在", str(context.exception))

    def test_get_document_text_invalid_format(self):
        """测试获取非docx文件的文本应该失败"""
        txt_file = os.path.join(self.temp_dir, "test.txt")
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write("这不是docx文件")

        with self.assertRaises(FileError) as context:
            get_document_text(txt_file)
        self.assertIn("文件格式不支持", str(context.exception))

    def test_get_document_text_empty_filepath(self):
        """测试空文件路径应该失败"""
        with self.assertRaises(ValidationError) as context:
            get_document_text(filepath="")
        self.assertIn("文件路径不能为空", str(context.exception))

    def test_get_document_text_return_fields(self):
        """测试返回结果包含所有必需的字段"""
        # 创建文档
        create_document(filepath=self.test_file)

        # 获取文档文本
        result = get_document_text(filepath=self.test_file)

        # 检查所有必需字段是否存在
        required_fields = [
            "message", "file_path", "text", "paragraph_count",
            "table_count", "total_characters", "total_words"
        ]

        for field in required_fields:
            self.assertIn(field, result, f"缺少字段: {field}")

    def test_get_document_text_complex_table(self):
        """测试复杂表格的文本提取"""
        create_document(filepath=self.test_file)

        doc = Document(self.test_file)

        # 创建复杂表格
        table = doc.add_table(rows=3, cols=3)
        test_data = [
            ["列1", "列2", "列3"],
            ["行1数据1", "行1数据2", "行1数据3"],
            ["行2数据1", "行2数据2", "行2数据3"]
        ]

        for i, row_data in enumerate(test_data):
            for j, cell_data in enumerate(row_data):
                table.cell(i, j).text = cell_data

        doc.save(self.test_file)

        # 获取文档文本
        result = get_document_text(filepath=self.test_file)

        # 验证表格内容被正确提取
        for row_data in test_data:
            for cell_data in row_data:
                self.assertIn(cell_data, result["text"])


if __name__ == '__main__':
    unittest.main()