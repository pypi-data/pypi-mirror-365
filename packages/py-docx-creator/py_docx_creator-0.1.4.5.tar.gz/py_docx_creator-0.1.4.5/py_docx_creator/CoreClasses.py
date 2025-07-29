from dataclasses import dataclass, fields
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.styles.styles import Styles
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from py_docx_creator.AbstractClasses import DocumentCreator, DocumentWriter, DocumentStyle, PageStyle, ParagraphStyle, \
    TextStyle, ParagraphStyleManager, TextStyleManager, PageStyleManager, AlignParagraph


# region _________________Управление документом_________________

class CoreDocumentCreator(DocumentCreator):

    def __init__(self):
        super().__init__()
        self._document = None
        self._file_name = None

    def create_document(self, file_name):
        self.document = Document()
        self.file_name = file_name

    def load_document(self):
        self.document = Document(self.path_to_document or self.file_name)

    def save_document(self):
        self.document.save(self.file_name)


class CoreDocumentWriter(DocumentWriter):

    @staticmethod
    def add_paragraph_to_document(document: Document) -> Paragraph:
        return document.add_paragraph()

    @staticmethod
    def add_run_to_paragraph(paragraph: Paragraph, text: str) -> Run:
        return paragraph.add_run(text)

    @staticmethod
    def add_page_break(document: Document) -> None:
        document.add_page_break()


class CoreDocumentStyle(DocumentStyle):

    def __init__(self):
        super().__init__()

    def get_document_style(self, document: Document) -> Styles:
        return document.style[f"{self.document_style}"]


# endregion _________________Управление документом_________________

# region _________________Стили_________________

@dataclass
class CoreTextStyle(TextStyle):
    """
        Стиль текста.

        Атрибуты:
            size ( float | None): # размер шрифта

            name ( str | None): # наименование шрифта

            bold ( bool | None): # жирное начертание шрифта

            italic ( bool | None): # курсивное начертание шрифта

            underline ( bool | None): # подчеркнутое начертание шрифта

        """

    size: float | None = None
    name: str | None = None
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None


@dataclass
class CorePageStyle(PageStyle):
    """
        Стиль страницы.

            Атрибуты:
                top_margin ( float | None): # отступ сверху

                bottom_margin ( float | None): # отступ снизу

                left_margin ( float | None): # отступ слева

                right_margin ( float | None): # отступ справа

        """

    top_margin: float | None = None
    bottom_margin: float | None = None
    left_margin: float | None = None
    right_margin: float | None = None


@dataclass
class CoreParagraphStyle(ParagraphStyle):
    """
    Стиль форматирования параграфа.

    Атрибуты:
        alignment (AlignParagraph | None): Выравнивание текста (влево, по центру, по ширине и т.п.).

        space_after (float | None): Отступ после параграфа.

        space_before (float | None): Отступ перед параграфом.

        left_indent (float | None): Отступ от левого края страницы.

        right_indent (float | None): Отступ от правого края страницы.

        line_spacing (float | None): Межстрочный интервал.

        first_line_indent (float | None): Отступ первой строки (красная строка).

        page_break_before (bool | None): Разрыв страницы перед параграфом.
    """
    alignment: AlignParagraph | None = None  # выравнивание
    space_after: float | None = None  # отступ до параграфа
    space_before: float | None = None  # отступ после параграфа
    left_indent: float | None = None  # отступ от левого края
    right_indent: float | None = None  # отступ от правого края
    line_spacing: float | None = None  # межстрочный интервал
    first_line_indent: float | None = None  # отступ красной строки
    page_break_before: bool | None = None  # разрыв страницы перед параграфом


# endregion _________________Стили_________________

# region _________________Управление стилями_________________

class CorePageStyleManager(PageStyleManager):
    @staticmethod
    def apply_style(document: Document, style: Any) -> None:
        """Применение стиля из dataclass к секциям документа."""

        for section in document.sections:
            for field in fields(style):
                value = getattr(style, field.name)
                if value is not None:

                    # проверка необходимости преобразования типа данных
                    if field.name in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
                        value = Pt(value)

                    setattr(section, field.name, value)


class CoreTextStyleManager(TextStyleManager):
    @staticmethod
    def apply_style(run: Run, style: Any) -> None:
        """Применение стиля"""
        for field in fields(style):
            value = getattr(style, field.name)
            if value is not None:

                # проверка необходимости преобразования типа данных
                if field.name in ("size",):
                    value = Pt(value)

                setattr(run.font, field.name, value)


class CoreParagraphStyleManager(ParagraphStyleManager):
    @staticmethod
    def apply_style(paragraph: Paragraph, style: Any) -> None:
        """Применение стиля к параграфу"""

        paragraph_style = paragraph.paragraph_format
        for field in fields(style):
            value = getattr(style, field.name)
            if value is not None:

                # проверка необходимости преобразования типа данных
                if field.name in ("space_after", "space_before", "left_indent", "right_indent", "first_line_indent"):
                    if field.name in ("left_indent", "right_indent"):
                        value = Inches(value)
                    else:
                        value = Pt(value)

                setattr(paragraph_style, field.name, value)


class CoreStyleManager:
    PAGE_STYLE_MANAGER: CorePageStyleManager = CorePageStyleManager
    TEXT_STYLE_MANAGER: CoreTextStyleManager = CoreTextStyleManager
    PARAGRAPH_STYLE_MANAGER: CoreParagraphStyleManager = CoreParagraphStyleManager

# endregion _________________Управление стилями_________________
