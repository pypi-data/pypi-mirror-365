"""Document conversion utilities with Marker integration and fallbacks."""

import logging
from pathlib import Path
from typing import Optional, Dict, Any
import hashlib
import json
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentConverter:
    """Converts various document formats to markdown with caching."""
    
    def __init__(self, cache_dir: Path, config: Optional[Dict[str, Any]] = None):
        self.cache_dir = cache_dir / "converted"
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        self.config = config or {}
        self.enabled = self.config.get("enabled", True)
        self.use_gpu = self.config.get("use_gpu", False)
        self.max_file_size_mb = self.config.get("max_file_size_mb", 100)
        self.supported_formats = self.config.get("formats", 
            ['pdf', 'docx', 'pptx', 'xlsx', 'epub'])
        
        # Lazy loading to avoid import errors if marker not installed
        self._marker_converter = None
        self._has_marker = None
        
    @property
    def has_marker(self) -> bool:
        """Check if Marker is available."""
        if self._has_marker is None:
            try:
                import marker
                self._has_marker = True
                logger.debug("Marker PDF converter available")
            except ImportError:
                self._has_marker = False
                logger.debug("Marker PDF converter not available, using fallbacks")
        return self._has_marker
    
    def _get_marker_converter(self):
        """Get Marker converter instance with lazy loading."""
        if not self.has_marker:
            return None
            
        if self._marker_converter is None:
            try:
                from marker.converters.pdf import PdfConverter
                from marker.models import create_model_dict
                
                # Create model dict with GPU/CPU preference
                model_dict = create_model_dict(device='cuda' if self.use_gpu else 'cpu')
                self._marker_converter = PdfConverter(artifact_dict=model_dict)
                logger.info(f"Initialized Marker converter (GPU: {self.use_gpu})")
            except Exception as e:
                logger.error(f"Failed to initialize Marker converter: {e}")
                self._marker_converter = None
                
        return self._marker_converter
    
    def _get_cache_path(self, file_path: Path) -> tuple[Path, Path]:
        """Get cache paths for converted content and metadata."""
        # Create hash of file path for unique cache key
        file_hash = hashlib.md5(str(file_path.absolute()).encode()).hexdigest()
        cache_name = f"{file_path.stem}_{file_hash}"
        
        content_path = self.cache_dir / f"{cache_name}.md"
        meta_path = self.cache_dir / f"{cache_name}.meta"
        
        return content_path, meta_path
    
    def _is_cached(self, file_path: Path) -> bool:
        """Check if file conversion is cached and up to date."""
        content_path, meta_path = self._get_cache_path(file_path)
        
        if not (content_path.exists() and meta_path.exists()):
            return False
            
        try:
            with open(meta_path, 'r') as f:
                meta = json.load(f)
            
            # Check if source file was modified since cache
            current_mtime = file_path.stat().st_mtime
            cached_mtime = meta.get('source_mtime', 0)
            
            return current_mtime <= cached_mtime
        except Exception as e:
            logger.debug(f"Cache check failed for {file_path}: {e}")
            return False
    
    def _save_to_cache(self, file_path: Path, content: str):
        """Save converted content to cache."""
        content_path, meta_path = self._get_cache_path(file_path)
        
        try:
            # Save content
            with open(content_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # Save metadata
            meta = {
                'source_path': str(file_path.absolute()),
                'source_mtime': file_path.stat().st_mtime,
                'converted_at': datetime.now().isoformat(),
                'converter': 'marker' if self.has_marker else 'fallback'
            }
            
            with open(meta_path, 'w') as f:
                json.dump(meta, f, indent=2)
                
            logger.debug(f"Cached conversion: {file_path.name}")
        except Exception as e:
            logger.error(f"Failed to save cache for {file_path}: {e}")
    
    def _load_from_cache(self, file_path: Path) -> Optional[str]:
        """Load converted content from cache."""
        content_path, _ = self._get_cache_path(file_path)
        
        try:
            with open(content_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.debug(f"Failed to load cache for {file_path}: {e}")
            return None
    
    def _convert_with_marker(self, file_path: Path) -> Optional[str]:
        """Convert document using Marker."""
        converter = self._get_marker_converter()
        if not converter:
            return None
            
        try:
            logger.info(f"Converting with Marker: {file_path.name}")
            
            # Convert document
            rendered = converter(str(file_path))
            
            # Extract text from rendered output
            from marker.output import text_from_rendered
            text, _, images = text_from_rendered(rendered)
            
            # TODO: Handle extracted images if needed
            if images:
                logger.debug(f"Extracted {len(images)} images from {file_path.name}")
            
            return text
            
        except Exception as e:
            logger.error(f"Marker conversion failed for {file_path}: {e}")
            return None
    
    def _convert_with_fallback(self, file_path: Path) -> Optional[str]:
        """Convert document using simple fallback methods."""
        ext = file_path.suffix.lower()
        
        try:
            if ext == '.pdf':
                return self._convert_pdf_fallback(file_path)
            elif ext == '.docx':
                return self._convert_docx_fallback(file_path)
            else:
                logger.warning(f"No fallback converter for {ext} files")
                return None
                
        except Exception as e:
            logger.error(f"Fallback conversion failed for {file_path}: {e}")
            return None
    
    def _convert_pdf_fallback(self, file_path: Path) -> Optional[str]:
        """Simple PDF text extraction using pypdf."""
        try:
            import pypdf
            
            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(f"## Page {page_num + 1}\n\n{text}\n")
                    except Exception as e:
                        logger.debug(f"Failed to extract page {page_num + 1}: {e}")
                        continue
            
            return "\n".join(text_parts) if text_parts else None
            
        except ImportError:
            logger.warning("pypdf not available for PDF fallback. Install with: pip install localgenius[documents]")
            return None
    
    def _convert_docx_fallback(self, file_path: Path) -> Optional[str]:
        """Simple DOCX text extraction using python-docx."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Simple markdown formatting
                    style = paragraph.style.name.lower()
                    if 'heading' in style:
                        level = 1
                        if 'heading 2' in style:
                            level = 2
                        elif 'heading 3' in style:
                            level = 3
                        text = f"{'#' * level} {text}"
                    
                    text_parts.append(text)
            
            # Add tables
            for table in doc.tables:
                table_md = self._table_to_markdown(table)
                if table_md:
                    text_parts.append(table_md)
            
            return "\n\n".join(text_parts) if text_parts else None
            
        except ImportError:
            logger.debug("python-docx not available for DOCX fallback")
            return None
    
    def _table_to_markdown(self, table) -> str:
        """Convert DOCX table to markdown format."""
        try:
            rows = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                
                # Add separator after header row
                if i == 0:
                    separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                    rows.append(separator)
            
            return "\n".join(rows)
        except Exception as e:
            logger.debug(f"Failed to convert table to markdown: {e}")
            return ""
    
    def convert_to_markdown(self, file_path: Path) -> Optional[str]:
        """Convert document to markdown with caching.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Markdown content or None if conversion failed
        """
        if not self.enabled:
            return None
            
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        # Check file size
        try:
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                logger.warning(f"File too large ({file_size_mb:.1f}MB): {file_path.name}")
                return None
        except Exception:
            pass
        
        # Check if format is supported
        ext = file_path.suffix.lower().lstrip('.')
        if ext not in self.supported_formats:
            logger.debug(f"Format not supported: {ext}")
            return None
        
        # Check cache first
        if self._is_cached(file_path):
            logger.debug(f"Loading from cache: {file_path.name}")
            cached = self._load_from_cache(file_path)
            if cached:
                return cached
        
        # Try conversion with Marker first, then fallback
        content = None
        
        if self.has_marker:
            content = self._convert_with_marker(file_path)
        
        if content is None:
            logger.debug(f"Trying fallback conversion for {file_path.name}")
            content = self._convert_with_fallback(file_path)
        
        # Cache successful conversion
        if content:
            self._save_to_cache(file_path, content)
            logger.info(f"Successfully converted: {file_path.name}")
        else:
            logger.warning(f"Failed to convert: {file_path.name}")
        
        return content
    
    def clear_cache(self):
        """Clear conversion cache."""
        try:
            import shutil
            if self.cache_dir.exists():
                shutil.rmtree(self.cache_dir)
                self.cache_dir.mkdir(parents=True, exist_ok=True)
                logger.info("Conversion cache cleared")
        except Exception as e:
            logger.error(f"Failed to clear cache: {e}")
    
    def get_cache_stats(self) -> Dict[str, Any]:
        """Get cache statistics."""
        try:
            cache_files = list(self.cache_dir.glob("*.md"))
            total_size = sum(f.stat().st_size for f in cache_files)
            
            return {
                "cache_dir": str(self.cache_dir),
                "cached_files": len(cache_files),
                "total_size_mb": total_size / (1024 * 1024),
                "marker_available": self.has_marker
            }
        except Exception as e:
            logger.error(f"Failed to get cache stats: {e}")
            return {"error": str(e)}