"""
Text extraction utilities for different file formats
"""

from pathlib import Path
from typing import Union

# --- Core Dependencies (already part of the main installation) ---
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

# --- Optional Imports for Extended Format Support ---
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    from odf.text import P
    from odf.teletype import extractText as odf_extract_text
    from odf.opendocument import load as load_odt
    ODT_AVAILABLE = True
except ImportError:
    ODT_AVAILABLE = False


class TextExtractorError(Exception):
    """Custom exception for text extraction errors."""
    pass


class TextExtractor:
    """A robust text extraction utility for various file formats."""

    @staticmethod
    def extract_text(file_path: Union[str, Path]) -> str:
        """
        Extracts text from a file, routing to the correct parser based on extension.
        Handles PDF, DOCX, TXT, images (JPG, PNG), HTML, and ODT.
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.stat().st_size == 0:
            raise TextExtractorError(f"File is empty: {file_path}")
        
        ext = file_path.suffix.lower()
        
        if ext == ".pdf":
            return TextExtractor._extract_pdf_text(file_path)
        elif ext == ".docx":
            return TextExtractor._extract_docx_text(file_path)
        elif ext == ".txt":
            return TextExtractor._extract_txt_text(file_path)
        elif ext in [".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".gif"]:
            return TextExtractor._extract_image_text_ocr(file_path)
        elif ext in [".html", ".htm"]:
            return TextExtractor._extract_html_text(file_path)
        elif ext == ".odt":
            return TextExtractor._extract_odt_text(file_path)
        else:
            raise TextExtractorError(f"Unsupported file type: '{ext}'. Supported types: PDF, DOCX, TXT, PNG, JPG, HTML, ODT.")

    @staticmethod
    def _extract_pdf_text(file_path: Path) -> str:
        """Optimized PDF text extraction, preferring PyMuPDF for speed and accuracy."""
        if PYMUPDF_AVAILABLE:
            try:
                with fitz.open(file_path) as doc:
                    text = "".join(page.get_text() for page in doc)
                if text.strip():
                    return text
            except Exception as e:
                # Fallback to pdfminer if PyMuPDF fails
                pass
        
        if PDFMINER_AVAILABLE:
            try:
                text = pdfminer_extract_text(str(file_path))
                if text.strip():
                    return text
            except Exception as e:
                raise TextExtractorError(f"PDFMiner extraction failed: {e}")

        raise TextExtractorError("No readable text found in PDF. The document may be an image-only PDF. Try installing with '[ocr]' support.")

    @staticmethod
    def _extract_docx_text(file_path: Path) -> str:
        """Extracts text from DOCX, including paragraphs and tables."""
        if not DOCX_AVAILABLE:
            raise TextExtractorError("python-docx is not installed. Please run 'pip install python-docx'.")
        
        try:
            doc = docx.Document(str(file_path))
            full_text = [p.text for p in doc.paragraphs]
            # Include table data
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            text = "\n".join(full_text)
            if not text.strip():
                raise TextExtractorError("No readable text found in DOCX file.")
            return text
        except Exception as e:
            raise TextExtractorError(f"DOCX extraction failed: {e}")

    @staticmethod
    def _extract_txt_text(file_path: Path) -> str:
        """Extracts text from TXT, trying multiple encodings."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    return f.read()
            except Exception as e:
                raise TextExtractorError(f"Failed to read text file with common encodings: {e}")
        except Exception as e:
            raise TextExtractorError(f"TXT extraction failed: {e}")

    @staticmethod
    def _extract_image_text_ocr(file_path: Path) -> str:
        """Extracts text from various image formats using Tesseract OCR."""
        if not OCR_AVAILABLE:
            raise TextExtractorError(
                "Image processing dependencies not installed. "
                "Run: pip install ai-resume-parser[ocr]"
            )
        try:
            text = pytesseract.image_to_string(Image.open(file_path))
            if not text.strip():
                raise TextExtractorError("OCR found no readable text in the image.")
            return text
        except pytesseract.TesseractNotFoundError:
            raise TextExtractorError(
                "Tesseract OCR engine not found. "
                "Please install it from https://github.com/tesseract-ocr/tesseract and ensure it's in your system's PATH."
            )
        except Exception as e:
            raise TextExtractorError(f"Image OCR extraction failed: {e}")

    @staticmethod
    def _extract_html_text(file_path: Path) -> str:
        """Extracts clean, readable text from HTML files."""
        if not HTML_AVAILABLE:
            raise TextExtractorError(
                "HTML parsing dependency not installed. "
                "Run: pip install ai-resume-parser[html]"
            )
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
            
            for tag in soup(["script", "style", "nav", "footer", "aside"]):
                tag.decompose()
            
            text = soup.get_text(separator="\n", strip=True)
            if not text.strip():
                raise TextExtractorError("No readable content found in HTML file.")
            return text
        except Exception as e:
            raise TextExtractorError(f"HTML extraction failed: {e}")

    @staticmethod
    def _extract_odt_text(file_path: Path) -> str:
        """Extracts text from OpenDocument Text (.odt) files."""
        if not ODT_AVAILABLE:
            raise TextExtractorError(
                "ODT parsing dependency not installed. "
                "Run: pip install ai-resume-parser[odt]"
            )
        try:
            doc = load_odt(str(file_path))
            paragraphs = doc.getElementsByType(P)
            text = "\n".join(odf_extract_text(p) for p in paragraphs)
            if not text.strip():
                raise TextExtractorError("No readable text found in ODT file.")
            return text
        except Exception as e:
            raise TextExtractorError(f"ODT extraction failed: {e}")
