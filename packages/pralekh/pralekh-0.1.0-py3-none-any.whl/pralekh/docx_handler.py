from docx import Document

def handle_docx(path, output_format="text"):
    doc = Document(path)
    text = "\n".join([para.text for para in doc.paragraphs])
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    return text if output_format == "text" else {"text": text, "tables": tables}
