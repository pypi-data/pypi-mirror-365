"""Main module."""

import pandas as pd
from matplotlib import pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import argparse
from datetime import datetime
import os


class ReportGenerator:
    def __init__(self, excel_file='budget.xlsx', output_file=None):
        self.excel_file = excel_file
        self.document = None
        self.data = None
        self._content_sections = None  # Cache for content.md sections

        # Create reports directory if it doesn't exist
        reports_dir = Path('reports')
        reports_dir.mkdir(exist_ok=True)

        # Simplified naming: project_report_yyyy-mm-dd.docx in reports folder
        if output_file is None:
            base_name = f'project_report_{datetime.now().strftime("%Y-%m-%d")}.docx'
            base_path = reports_dir / base_name
            self.output_file = str(self._get_unique_filename(base_path))
        else:
            # If custom filename provided, still put it in reports folder
            self.output_file = str(reports_dir / output_file)

    def _get_unique_filename(self, base_filepath):
        """Ensure unique filename by adding increment if file exists"""
        if not base_filepath.exists():
            return base_filepath

        # If file exists, add increment: project_report_2025-01-25_v2.docx
        name_stem = base_filepath.stem  # project_report_2025-01-25
        extension = base_filepath.suffix  # .docx
        parent_dir = base_filepath.parent

        counter = 2
        while True:
            new_filename = f"{name_stem}_v{counter}{extension}"
            new_filepath = parent_dir / new_filename
            if not new_filepath.exists():
                print(f"📝 File exists, creating new version: {new_filename}")
                return new_filepath
            counter += 1

            # Safety break to avoid infinite loop
            if counter > 100:
                # Use timestamp as fallback
                timestamp = datetime.now().strftime("%H%M%S")
                return parent_dir / f"{name_stem}_{timestamp}{extension}"

    def _load_content_sections(self):
        """Load and cache all sections from content.md file once"""
        if self._content_sections is not None:
            return self._content_sections

        try:
            content_file = Path('content.md')
            if not content_file.exists():
                print(f"Info: content.md not found, using default content")
                self._content_sections = {}
                return self._content_sections

            with open(content_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Split content by headers (# Section Name)
            sections = {}
            current_section = None
            current_content = []

            for line in content.split('\n'):
                if line.startswith('# '):
                    # Save previous section if exists
                    if current_section:
                        sections[current_section] = '\n'.join(
                            current_content).strip()
                    # Start new section
                    current_section = line[2:].strip()  # Remove '# '
                    current_content = []
                else:
                    current_content.append(line)

            # Save last section
            if current_section:
                sections[current_section] = '\n'.join(current_content).strip()

            self._content_sections = sections
            return self._content_sections

        except Exception as e:
            print(f"Error reading content.md: {e}")
            self._content_sections = {}
            return self._content_sections

    def read_section_from_content(self, section_name):
        """Read a specific section from cached content.md sections"""
        sections = self._load_content_sections()
        return sections.get(section_name, None)

    def add_markdown_content(self, section_name, default_content=None):
        """Add content from content.md section to document with basic formatting"""
        content = self.read_section_from_content(section_name)

        if not content and default_content:
            content = default_content
        elif not content:
            return

        # Split content into paragraphs
        paragraphs = content.split('\n\n')

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Handle bullet points (markdown style) - removed bold formatting
            if para.startswith('- ') or para.startswith('* '):
                # Split multiple bullet points
                bullets = [line.strip()[2:] for line in para.split(
                    '\n') if line.strip().startswith(('- ', '* '))]
                for bullet in bullets:
                    self.document.add_paragraph(bullet, style='List Bullet')
            else:
                # Regular paragraph
                self.document.add_paragraph(para)

    def _remove_table_borders(self, table):
        """Efficiently remove borders from table"""
        # Set table style to None to remove default borders
        table.style = None

        # Remove borders at XML level for all existing and new cells
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # Remove borders if they exist
                tcBorders = tcPr.find(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
                if tcBorders is not None:
                    tcPr.remove(tcBorders)

    def _format_cell_alignment(self, cell, column_index, is_header=False, is_total_row=False):
        """Efficiently format cell alignment and styling"""
        for paragraph in cell.paragraphs:
            # Set font size
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(11)

            # Right-align numeric columns (1, 2, 3)
            if column_index in [1, 2, 3]:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Bold formatting for headers and totals
            if is_header or is_total_row:
                for run in paragraph.runs:
                    run.bold = True

    def load_data(self):
        """Load budget data from Excel file"""
        try:
            self.data = pd.read_excel(self.excel_file)
            print(f"✅ Successfully loaded data from {self.excel_file}")
            print(f"   Columns: {list(self.data.columns)}")
            print(f"   Rows: {len(self.data)}")
            return True
        except FileNotFoundError:
            print(f"❌ Error: Could not find {self.excel_file}")
            return False
        except Exception as e:
            print(f"❌ Error loading data: {e}")
            return False

    def create_document(self):
        """Initialize the Word document"""
        self.document = Document()
        # Add title with current date
        title = f'Budget Report - {datetime.now().strftime("%B %d, %Y")}'
        self.document.add_heading(title, 0)
        print("✅ Document initialized")

    def add_introduction(self):
        """Add introduction section"""
        self.document.add_heading('Summary', level=1)

        default_intro = f"""This report provides a comprehensive overview of budget allocation and expenditure status as of {datetime.now().strftime("%B %d, %Y")}.

Key metrics include budget utilization rates, remaining fund allocation, and project-specific financial performance indicators."""

        self.add_markdown_content('Summary', default_intro)
        print("✅ Summary section added")

    def add_deliverables_progress(self):
        """Add deliverables progress section"""
        self.document.add_heading('Deliverables Progress', level=1)

        default_deliverables = """Progress on key project deliverables remains on track with established timelines.

- Major milestones achieved during this reporting period
- Current status of ongoing deliverables
- Any adjustments to delivery schedules"""

        self.add_markdown_content(
            'Deliverables Progress', default_deliverables)
        print("✅ Deliverables Progress section added")

    def add_budget_table(self):
        """Add budget data table"""
        if self.data is None:
            print("❌ No data available for table")
            return False

        self.document.add_heading('Budget', level=1)

        # Add summary paragraph
        total_budgeted = self.data['Budgeted'].sum(
        ) if 'Budgeted' in self.data.columns else 0
        total_remaining = self.data['Remaining'].sum(
        ) if 'Remaining' in self.data.columns else 0
        utilization_rate = ((total_budgeted - total_remaining) /
                            total_budgeted * 100) if total_budgeted > 0 else 0

        summary_text = f"Total Budget: ${total_budgeted:,.0f} | Utilization Rate: {utilization_rate:.1f}% | Remaining: ${total_remaining:,.0f}"
        self.document.add_paragraph(summary_text)

        # Create table
        table = self.document.add_table(rows=1, cols=len(self.data.columns))

        # Remove all borders efficiently
        self._remove_table_borders(table)

        # Add header row with formatting
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(self.data.columns):
            hdr_cells[i].text = str(column_name)
            self._format_cell_alignment(hdr_cells[i], i, is_header=True)

        # Add data rows efficiently
        for i, row_data in self.data.iterrows():
            row_cells = table.add_row().cells

            # Remove borders from new row
            for cell in row_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
                if tcBorders is not None:
                    tcPr.remove(tcBorders)

            is_total_row = 'TOTAL' in str(row_data.iloc[0]).upper()

            for j, cell_value in enumerate(row_data):
                # Format numbers with commas if numeric
                if pd.api.types.is_numeric_dtype(type(cell_value)) and pd.notna(cell_value):
                    row_cells[j].text = f"{cell_value:,.0f}" if cell_value == int(
                        cell_value) else f"{cell_value:,.2f}"
                else:
                    row_cells[j].text = str(cell_value)

                # Apply formatting
                self._format_cell_alignment(
                    row_cells[j], j, is_total_row=is_total_row)

        print("✅ Budget table added")
        return True

    def add_key_points(self):
        """Add key points section"""
        # Add space and intro sentence
        self.document.add_paragraph("")  # Empty paragraph for spacing
        self.document.add_paragraph("Summary of budget status:")

        # Generate dynamic key points based on data
        key_points = []

        if self.data is not None and all(col in self.data.columns for col in ['Budgeted', 'Remaining']):
            # Calculate insights efficiently
            total_budgeted = self.data['Budgeted'].sum()
            total_remaining = self.data['Remaining'].sum()
            utilization_rate = ((total_budgeted - total_remaining) /
                                total_budgeted * 100) if total_budgeted > 0 else 0

            # Find highest and lowest utilization tasks efficiently
            self.data['Utilization%'] = (
                (self.data['Budgeted'] - self.data['Remaining']) / self.data['Budgeted'] * 100).round(1)

            if len(self.data) > 1:
                # Filter out totals row efficiently
                non_total_data = self.data[~self.data['Task'].str.contains(
                    'TOTAL', case=False, na=False)] if 'Task' in self.data.columns else self.data

                if len(non_total_data) > 0:
                    highest_util_idx = non_total_data['Utilization%'].idxmax()
                    lowest_util_idx = non_total_data['Utilization%'].idxmin()

                    highest_util = non_total_data.loc[highest_util_idx]
                    lowest_util = non_total_data.loc[lowest_util_idx]

                    key_points = [
                        f"Overall budget utilization stands at {utilization_rate:.1f}%",
                        f"Highest utilization: {highest_util['Task']} at {highest_util['Utilization%']:.1f}%",
                        f"Lowest utilization: {lowest_util['Task']} at {lowest_util['Utilization%']:.1f}%",
                        f"Total remaining funds: ${total_remaining:,.0f}"
                    ]

        # Use default if no dynamic points generated
        if not key_points:
            key_points = [
                "Budget tracking is current and accurate",
                "All expenditures are within approved parameters",
                "Financial controls are operating effectively",
                "Regular monitoring continues as scheduled"
            ]

        default_key_points = "\n".join([f"- {point}" for point in key_points])
        self.add_markdown_content('key_points', default_key_points)

        print("✅ Key points section added")

    def add_budget_chart(self):
        """Add budget visualization chart"""
        if self.data is None:
            print("❌ No data available for chart")
            return False

        # Add chart description
        chart_desc = "Figure 1 provides a visual comparison of budgeted amounts versus remaining funds for each project component."
        self.add_markdown_content('chart_description', chart_desc)

        try:
            # Set matplotlib to non-interactive backend to prevent chart from showing
            plt.ioff()

            # Filter out TOTALS row for better visualization
            chart_data = self.data[~self.data['Task'].str.contains(
                'TOTAL', case=False, na=False)] if 'Task' in self.data.columns else self.data

            if len(chart_data) == 0:
                chart_data = self.data

            # Create chart more efficiently
            fig, ax = plt.subplots(figsize=(12, 8))

            # Plot data
            x_pos = range(len(chart_data))
            ax.bar([x - 0.2 for x in x_pos], chart_data['Budgeted'],
                   0.4, label='Budgeted', color='#2E8B57')
            ax.bar([x + 0.2 for x in x_pos], chart_data['Remaining'],
                   0.4, label='Remaining', color='#4169E1')

            # Formatting
            ax.set_title('Budget Status by Task', fontsize=16,
                         fontweight='bold', pad=20)
            ax.set_xlabel('Project Tasks', fontsize=12)
            ax.set_ylabel('Amount ($)', fontsize=12)
            ax.set_xticks(x_pos)
            ax.set_xticklabels(chart_data['Task'] if 'Task' in chart_data.columns else chart_data.index,
                               rotation=45, ha='right')
            ax.legend(loc='upper right')
            ax.yaxis.set_major_formatter(
                plt.FuncFormatter(lambda x, p: f'{x:,.0f}'))
            ax.grid(axis='y', alpha=0.3)

            plt.tight_layout()

            # Save chart
            chart_filename = 'budget_chart.png'
            plt.savefig(chart_filename, bbox_inches='tight',
                        dpi=300, facecolor='white')

            # Add to document
            self.document.add_picture(chart_filename, width=Inches(6.5))

            # Add figure caption in 9pt font
            caption_paragraph = self.document.add_paragraph(
                "Figure 1: Total amounts budgeted and remaining by project task.")
            caption_run = caption_paragraph.runs[0]
            caption_run.font.size = Pt(9)

            # Close the figure to prevent display and free memory
            plt.close(fig)

            print("✅ Budget chart added")
            return True

        except Exception as e:
            print(f"❌ Error creating chart: {e}")
            return False

    def add_challenges(self):
        """Add challenges section"""
        self.document.add_heading('Challenges', level=1)

        default_challenges = """Current challenges and mitigation strategies:

- Resource allocation constraints and proposed solutions
- Technical obstacles encountered and resolution approaches
- Timeline adjustments required due to external factors"""

        self.add_markdown_content('Challenges', default_challenges)
        print("✅ Challenges section added")

    def add_next_period_activities(self):
        """Add next period activities section"""
        self.document.add_heading('Next Period Activities', level=1)

        default_next_period = """Planned activities for the upcoming reporting period:

- Priority tasks and deliverables for next phase
- Resource requirements and allocation plans
- Key milestones and target completion dates"""

        self.add_markdown_content(
            'Next Period Activities', default_next_period)
        print("✅ Next Period Activities section added")

    def save_document(self):
        """Save the Word document"""
        try:
            self.document.save(self.output_file)
            print(f"✅ Report saved successfully: {self.output_file}")
            return True
        except Exception as e:
            print(f"❌ Error saving document: {e}")
            return False

    def rpt_pdf(self, word_file=None):
        """Convert Word report to PDF with the same name"""
        # Import pdf module - handle both package and script execution contexts
        try:
            from .pdf import convert_to_pdf
        except ImportError:
            # If relative import fails, try absolute import for script execution
            import sys
            from pathlib import Path
            current_dir = Path(__file__).parent
            sys.path.insert(0, str(current_dir))
            from pdf import convert_to_pdf

        # Use the current output file if no word_file specified
        if word_file is None:
            word_file = self.output_file

        success, result = convert_to_pdf(word_file)
        return success

    def generate_report(self):
        """Generate the complete report"""
        print("🚀 Starting report generation...")

        # Load data
        if not self.load_data():
            return False

        # Create document
        self.create_document()

        # Add sections
        self.add_introduction()
        self.add_deliverables_progress()
        self.add_budget_table()
        self.add_key_points()
        self.add_budget_chart()
        self.add_challenges()
        self.add_next_period_activities()

        # Save document
        success = self.save_document()

        if success:
            print(f"🎉 Report generation completed successfully!")
            print(f"📄 Output file: {self.output_file}")

            # Print file size
            file_size = os.path.getsize(self.output_file) / 1024  # KB
            print(f"📊 File size: {file_size:.1f} KB")

        return success


def main():
    """Main function with command line argument support"""
    parser = argparse.ArgumentParser(
        description='Generate automated budget report',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python autorpt.py                               # Generate Word report
  python autorpt.py --pdf                         # Generate Word + PDF
  python autorpt.py --pdf-only                    # Convert existing reports to PDF
  python autorpt.py --pdf-all                     # Convert all reports to PDF
  python autorpt.py -i budget.xlsx --pdf          # Custom input with PDF
        """)

    parser.add_argument('--input', '-i', default='budget.xlsx',
                        help='Input Excel file (default: budget.xlsx)')
    parser.add_argument('--output', '-o',
                        help='Output Word document filename')
    parser.add_argument('--pdf', '-p', action='store_true',
                        help='Also convert the report to PDF')
    parser.add_argument('--pdf-only', action='store_true',
                        help='Only convert existing Word reports to PDF (no new report generation)')
    parser.add_argument('--pdf-all', action='store_true',
                        help='Convert all Word reports in reports/ folder to PDF')
    parser.add_argument('--verbose', '-v', action='store_true',
                        help='Enable verbose output')

    args = parser.parse_args()

    if args.verbose:
        print("🔧 Verbose mode enabled")
        print(f"📁 Input file: {args.input}")
        print(f"📄 Output file: {args.output or 'auto-generated'}")
        if args.pdf:
            print("📄 PDF conversion enabled")
        if args.pdf_only:
            print("📄 PDF-only mode (no report generation)")
        if args.pdf_all:
            print("📄 Converting all reports to PDF")

    # Handle PDF-only operations
    if args.pdf_only or args.pdf_all:
        # Import pdf module - handle both package and script execution contexts
        try:
            from .pdf import convert_to_pdf, convert_all_reports
        except ImportError:
            # If relative import fails, try absolute import for script execution
            import sys
            from pathlib import Path
            current_dir = Path(__file__).parent
            sys.path.insert(0, str(current_dir))
            from pdf import convert_to_pdf, convert_all_reports

        if args.pdf_all:
            print("📁 Converting all Word reports to PDF...")
            results = convert_all_reports("reports")
            success = results["failed"] == 0
            if success:
                print(
                    f"🎉 Successfully converted {results['success']} report(s) to PDF!")
            else:
                print(
                    f"⚠️  Converted {results['success']} report(s), but {results['failed']} failed")
        elif args.pdf_only:
            # Convert the most recent report or specified output file
            if args.output:
                report_file = f"reports/{args.output}"
            else:
                # Find the most recent report
                from pathlib import Path
                from datetime import datetime
                reports_dir = Path('reports')
                if reports_dir.exists():
                    docx_files = sorted(reports_dir.glob(
                        "*.docx"), key=lambda x: x.stat().st_mtime, reverse=True)
                    if docx_files:
                        report_file = str(docx_files[0])
                        print(
                            f"🔍 Converting most recent report: {docx_files[0].name}")
                    else:
                        print("❌ No Word reports found in reports/ folder")
                        return 1
                else:
                    print("❌ Reports directory not found")
                    return 1

            success, result = convert_to_pdf(report_file)
            if success:
                print(f"🎉 PDF conversion completed successfully!")
            else:
                print(f"❌ PDF conversion failed: {result}")

        return 0 if success else 1

    # Normal report generation flow
    generator = ReportGenerator(args.input, args.output)
    success = generator.generate_report()

    # Convert to PDF if requested and report generation was successful
    if success and args.pdf:
        pdf_success = generator.rpt_pdf()
        if not pdf_success:
            print("⚠️  Report generated successfully but PDF conversion failed")

    return 0 if success else 1


def convert_to_pdf(word_file):
    """Standalone function to convert a Word document to PDF

    Args:
        word_file (str): Path to the Word document to convert

    Returns:
        bool: True if conversion successful, False otherwise
    """
    # Import pdf module - handle both package and script execution contexts
    try:
        from .pdf import convert_to_pdf as pdf_convert
    except ImportError:
        # If relative import fails, try absolute import for script execution
        import sys
        from pathlib import Path
        current_dir = Path(__file__).parent
        sys.path.insert(0, str(current_dir))
        from pdf import convert_to_pdf as pdf_convert

    success, result = pdf_convert(word_file)
    return success


if __name__ == "__main__":
    exit(main())
