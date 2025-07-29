from flask import Flask, jsonify, send_file
from docx import Document
from io import BytesIO
import uuid
import time
import threading
from flask_cors import CORS

app = Flask(__name__)
CORS(app, resources=r'/*')
# 内存存储文件（生产环境请使用数据库或Redis）
file_store = {}

@app.route('/generate_doc')
def generate_doc():
    """生成Word文档并返回下载URL"""
    try:
        # 创建内存中的Word文档
        buffer = BytesIO()
        doc = Document()
        doc.add_paragraph('这是自动生成的文档内容')
        doc.add_heading('文档标题', level=1)
        doc.add_paragraph(f'文档ID: {uuid.uuid4().hex}')
        doc.add_paragraph(f'生成时间: {time.strftime("%Y-%m-%d %H:%M:%S")}')
        doc.save(buffer)
        buffer.seek(0)  # 重置文件指针到开头
        
        # 生成文件ID和下载URL
        file_id = uuid.uuid4().hex
        download_url = f"/download/{file_id}"
        
        # 存储文件内容
        file_store[file_id] = {
            'content': buffer.read(),
            'filename': f"document_{file_id}.docx",
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'created_at': time.time()
        }
        
        # 返回JSON响应
        return jsonify({
            'success': True,
            'message': '文档生成成功',
            'download_url': download_url,
            'filename': file_store[file_id]['filename'],
            'expires_in': 300  # 5分钟过期
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/download/<file_id>')
def download_file(file_id):
    """通过文件ID下载文档"""
    if file_id not in file_store:
        return jsonify({
            'success': False,
            'error': f'文件不存在或已过期 (ID: {file_id})'
        }), 404
        
    file_info = file_store[file_id]
    
    # 检查文件是否过期（5分钟）
    if time.time() - file_info['created_at'] > 300:
        # 清理过期文件
        del file_store[file_id]
        return jsonify({
            'success': False,
            'error': '下载链接已过期'
        }), 410  # Gone
    
    # 创建文件流
    buffer = BytesIO(file_info['content'])
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=file_info['filename'],
        mimetype=file_info['mimetype']
    )

# 清理线程函数
def cleanup_expired_files():
    """定期清理过期的文件"""
    while True:
        time.sleep(60)  # 每分钟清理一次
        current_time = time.time()
        expired_ids = []
        for file_id, file_info in file_store.items():
            if current_time - file_info['created_at'] > 300:  # 超过5分钟
                expired_ids.append(file_id)
        
        for file_id in expired_ids:
            del file_store[file_id]
        
        if expired_ids:
            print(f"[清理线程] 清理了 {len(expired_ids)} 个过期文件")

# 启动清理线程
if __name__ == '__main__':
    # 启动清理线程
    threading.Thread(target=cleanup_expired_files, daemon=True).start()
    
    # 运行Flask应用
    print("启动Word文档生成服务...")
    print(f"访问 http://localhost:5000 或 http://localhost:5000/generate_doc")
    app.run(port=5000, debug=True)