from flask import Flask, jsonify, send_file
from docx import Document
from io import BytesIO
import uuid
import os
from flask_cors import CORS

app = Flask(__name__)
CORS(app, resources=r'/*')
# 存储内存中的文件（简单实现，生产环境应使用数据库或Redis）
file_store = {}

# @app.route('/generate_doc')
def generate_doc():
    """生成Word文档并返回下载URL"""
    try:
        # 1. 在内存中创建Word文档
        buffer = BytesIO()
        doc = Document()
        doc.add_paragraph("123")
        doc.add_heading('文档标题', level=1)
        doc.add_paragraph(f'文档ID: {uuid.uuid4().hex}')
        doc.save(buffer)
        buffer.seek(0)  # 重置指针到文件开头
        
        # 2. 生成唯一文件ID和下载URL
        file_id = uuid.uuid4().hex
        download_url = f"http://localhost:5000/download/{file_id}"
        
        # 3. 将文件内容存储在内存字典中
        file_store[file_id] = {
            'content': buffer.read(),
            'filename': f"document_{file_id}.docx",
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        print(download_url)
        # 4. 返回下载URL
        return 0
    
    except Exception as e:
        app.logger.error(f"生成文档失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': '文档生成失败',
            'details': str(e)
        }), 500

@app.route('/download/<file_id>')
def download_file(file_id):
    """通过文件ID提供下载"""
    if file_id not in file_store:
        return jsonify({
            'success': False,
            'error': '文件不存在或已过期'
        }), 404
    
    try:
        # 1. 从内存存储中获取文件信息
        file_info = file_store[file_id]
        
        # 2. 创建内存文件流
        buffer = BytesIO(file_info['content'])
        
        # 3. 返回文件下载
        return send_file(
            buffer,
            as_attachment=True,
            download_name=file_info['filename'],
            mimetype=file_info['mimetype']
        )
    
    except Exception as e:
        app.logger.error(f"下载文件失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': '文件下载失败',
            'details': str(e)
        }), 500

# 可选：定时清理过期文件
import threading
import time

def cleanup_expired_files():
    """定期清理过期的内存文件"""
    while True:
        time.sleep(60)  # 每分钟检查一次
        # 这里可以添加清理逻辑（如基于创建时间）
        # 示例：删除超过10分钟的文件
        current_time = time.time()
        expired_ids = [id for id, file in file_store.items() 
                       if current_time - file['created_at'] > 600]
        for id in expired_ids:
            del file_store[id]
        app.logger.info(f"清理了 {len(expired_ids)} 个过期文件")

# 启动清理线程
if __name__ == '__main__':
    # 在生产环境中应使用更可靠的任务调度
    cleanup_thread = threading.Thread(target=cleanup_expired_files, daemon=True)
    cleanup_thread.start()
    generate_doc()
    app.run(host="localhost", port=5000, debug=True)