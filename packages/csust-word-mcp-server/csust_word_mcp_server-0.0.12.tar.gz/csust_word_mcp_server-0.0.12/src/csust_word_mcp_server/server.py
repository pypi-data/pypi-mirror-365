from fastmcp import FastMCP
from docx import Document 
from io import BytesIO
from flask import Flask, jsonify, send_file
import uuid
import threading
import time
from flask_cors import CORS

API_BASE = "https://word-assistant"

mcp = FastMCP('word-assistant')
app = Flask(__name__)
CORS(app, resources=r'/*')

file_store = {}

@mcp.tool(name='word文档生成助手')
async def generate(text: str) -> str:
    """word生成

    Args:
        text: 输入的文本内容
    """
    """生成Word文档并返回下载URL"""
    try:
        # 1. 在内存中创建Word文档
        buffer = BytesIO()
        doc = Document()
        doc.add_paragraph(text)
        doc.add_heading('文档标题', level=1)
        doc.add_paragraph(f'文档ID: {uuid.uuid4().hex}')
        doc.save(buffer)
        buffer.seek(0)  # 重置指针到文件开头
        
        # 2. 生成唯一文件ID和下载URL
        file_id = uuid.uuid4().hex
        download_url = f"http://localhost:5000/download/{file_id}"
        
        # 3. 将文件内容存储在内存字典中
        file_store[file_id] = {
            'content': buffer.read(),
            'filename': f"document_{file_id}.docx",
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
    except Exception as e:
        app.logger.error(f"生成文档失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': '文档生成失败',
            'details': str(e)
        }), 500

    return download_url

@app.route('/download/<file_id>')
def download_file(file_id):
    """通过文件ID提供下载"""
    if file_id not in file_store:
        return jsonify({
            'success': False,
            'error': '文件不存在或已过期'
        }), 404
    
    try:
        # 1. 从内存存储中获取文件信息
        file_info = file_store[file_id]
        
        # 2. 创建内存文件流
        buffer = BytesIO(file_info['content'])
        
        # 3. 返回文件下载
        return send_file(
            buffer,
            as_attachment=True,
            download_name=file_info['filename'],
            mimetype=file_info['mimetype']
        )
    
    except Exception as e:
        app.logger.error(f"下载文件失败: {str(e)}")
        return jsonify({
            'success': False,
            'error': '文件下载失败',
            'details': str(e)
        }), 500

def cleanup_expired_files():
    """定期清理过期的内存文件"""
    while True:
        time.sleep(60)  # 每分钟检查一次
        # 删除超过10分钟的文件
        current_time = time.time()
        expired_ids = [id for id, file in file_store.items() 
                       if current_time - file['created_at'] > 600]
        for id in expired_ids:
            del file_store[id]
        app.logger.info(f"清理了 {len(expired_ids)} 个过期文件")

def app_run():
    app.run(host="localhost", port=5000)   

def mcp_run():
    mcp.run(transport="stdio")    

def run():
    cleanup_thread = threading.Thread(target=cleanup_expired_files, daemon=True)
    app_thread = threading.Thread(target=app_run)
    mcp_thread = threading.Thread(target=mcp_run)
    cleanup_thread.start()
    app_thread.start()
    mcp_thread.start()

if __name__ == "__main__":
    run()
    # asyncio.run(mcp.run_sse_async(host="0.0.0.0", port=8000))



    
