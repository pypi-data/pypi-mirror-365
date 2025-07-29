import re
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def parse_structured_text(text):
    """解析结构化公文文本（顺序敏感版）"""
    # 定义所有可能的字段
    fields = [
        "发文机关标志：", "发文字号：", "正文标题：", 
        "一级标题：", "二级标题：", "三级标题：", 
        "正文：", "发文机关：", "成文日期："
    ]
    
    # 创建正则表达式模式
    pattern = r"(" + "|".join(re.escape(field) for field in fields) + r")" + r"([\s\S]*?)(?=\n\n(?:发文机关标志：|发文字号：|正文标题：|一级标题：|二级标题：|三级标题：|正文：|发文机关：|成文日期：)|\Z)"
    pattern = re.compile(pattern, re.DOTALL)
    
    # 查找所有匹配项
    matches = pattern.findall(text)
    
    # 处理结果（保留顺序）
    result = []
    for field, content in matches:
        content = content.strip()
        result.append((field, content))
    
    return result

def create_word_document(parsed_data):
    """根据解析数据创建Word文档（顺序敏感版）"""
    doc = Document()
    
    # 设置文档默认字体
    doc.styles['Normal'].font.name = '仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    doc.styles['Normal'].font.size = Pt(16)  # 三号
    
    # 按原始顺序处理所有字段
    for field, content in parsed_data:
        # 1. 发文机关标志（居中、加粗、大字号、红色）
        if field == "发文机关标志：":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.bold = True
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
            run.font.size = Pt(22)  # 二
            run.font.color.rgb = RGBColor(255, 0, 0)
        
        # 2. 发文字号（居中）
        elif field == "发文字号：":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.font.size = Pt(16)  # 三号
            # #分割线
            # line = doc.add_paragraph()
            # run = line.add_run('-'*10)
            # run.font.color.rgb = RGBColor(255, 0, 0)
        
        # 3. 正文标题（居中、加粗）
        elif field == "正文标题：":
            doc.add_paragraph()  # 空行
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(18)  # 小二加粗
        
        # 4. 一级标题（加粗）
        elif field == "一级标题：":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(16)
        
        # 5. 二级标题（加粗）
        elif field == "二级标题：":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(15)
        
        # 6. 三级标题（加粗）
        elif field == "三级标题：":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(14)
        
        # 7. 正文（首行缩进）
        elif field == "正文：":
            # 处理正文中的多段落
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():  # 忽略空行
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Inches(0.4)  # 首行缩进2字符
                    run = p.add_run(para.strip())
            
        
        # 8. 发文机关（右对齐）
        elif field == "发文机关：":
            doc.add_paragraph()  # 空行
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(content)
            run.bold = True
        
        # 9. 成文日期（右对齐）
        elif field == "成文日期：":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(content)
    
    # 返回文档
    return doc

# 示例使用
if __name__ == "__main__":
    # 新的JSON数据
    json_data = '''
    {
        "output": "发文机关标志：\\n长沙理工大学文件\\n\\n发文字号：\\n长理工保〔202X〕XX号\\n\\n正文标题：\\n关于做好暑假期间校园内消防安全工作的通知\\n\\n\\n正文：\\n为切实加强我校暑假期间校园消防安全管理，有效防范和遏制各类火灾事故的发生，确保师生员工生命财产安全及校园稳定，现就有关工作通知如下：\\n\\n一级标题：\\n一、提高思想认识，强化责任担当\\n\\n正文：\\n各二级学院、职能部门要充分认识到暑期消防安全工作的特殊性和重要性。因假期留校师生相对集中，实验室设备持续运行，加之高温天气影响，电气火灾风险显著增加。各单位主要负责人作为第一责任人，必须亲自部署、靠前指挥，严格执行\\"党政同责、一岗双责\\"要求，将消防责任细化到岗、落实到人。\\n\\n一级标题：\\n二、开展隐患排查，落实整改措施\\n\\n正文：\\n（一）全面自查自纠。8月1日前完成首轮排查，重点检查学生宿舍违规电器使用、实验室危化品存储规范、配电房线路负荷情况、消防设施完好率等关键环节。建立隐患台账并实行销号管理，对发现的私拉乱接电线、堵塞疏散通道等问题立即整改。\\n（二）加强重点管控。对机电厂、图书馆古籍库房等高风险场所实施每日巡查，严格审批动火作业许可；强化施工工地临时用电监管，严禁无证焊接切割操作。\\n（三）完善应急准备。检验消防报警系统联动功能，补充更换过期灭火器材，组织保安队伍开展夜间应急拉动演练。\\n\\n一级标题：\\n三、深化宣传教育，提升防范能力\\n\\n正文：\\n通过校园网、微信公众号推送典型案例警示片，制作发放《暑期防火手册》。面向留校学生开展灭火实操培训，特别要加强对科研团队的新进人员进行危化品安全处置指导。在重点区域增设警示标识，利用电子屏滚动播放防火标语。\\n\\n一级标题：\\n四、严格值班值守，畅通信息渠道\\n\\n正文：\\n执行领导带班和24小时专人值班制度，安保部门每日进行防火巡查并做好记录。一旦发生火情，须按预案迅速启动应急响应，第一时间上报校总值班室（电话：XXXXXXX）及辖区消防大队。重大险情要同步向主管校领导汇报。\\n\\n发文机关：\\n长沙理工大学\\n\\n成文日期：\\n202X年XX月XX日"
    }
    '''
    
    # 解析JSON
    data = json.loads(json_data)
    output_text = data["output"]
    
    # 解析结构化文本（保留顺序）
    parsed_data = parse_structured_text(output_text)
    
    # 创建Word文档
    doc = create_word_document(parsed_data)